"""
rag_pipeline.py — Text Extraction + Chunking Engine

Handles:
  - PDF / DOCX / TXT / MD text extraction
  - Semantic chunking with overlap
  - Batch processing for embedding pipeline
"""
import io
import re
import logging
from typing import List, Dict, Any, Optional
from dataclasses import dataclass, field

logger = logging.getLogger(__name__)

# ═══════════════════════════════════════════════════════════
# Data Types
# ═══════════════════════════════════════════════════════════

@dataclass
class Chunk:
    chunk_id: str
    file_id: str
    content: str
    index: int
    metadata: Dict[str, Any] = field(default_factory=dict)


# ═══════════════════════════════════════════════════════════
# Text Extraction
# ═══════════════════════════════════════════════════════════

def extract_pdf(content: bytes) -> str:
    """Extract text from PDF bytes."""
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(io.BytesIO(content))
        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        return "\n\n".join(text_parts)
    except Exception as e:
        logger.error(f"PDF extraction failed: {e}")
        raise ValueError(f"PDF 解析失败: {e}")


def extract_docx(content: bytes) -> str:
    """Extract text from DOCX bytes."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(content))
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                if any(row_text):
                    paragraphs.append(" | ".join(row_text))
        return "\n\n".join(paragraphs)
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        raise ValueError(f"DOCX 解析失败: {e}")


def extract_text(content: bytes) -> str:
    """Extract text from TXT/MD bytes."""
    try:
        return content.decode("utf-8")
    except UnicodeDecodeError:
        try:
            return content.decode("latin-1")
        except Exception as e:
            raise ValueError(f"文本解码失败: {e}")


def extract_text_from_file(content: bytes, file_type: str) -> str:
    """Route to appropriate extractor based on file type."""
    extractors = {
        "pdf": extract_pdf,
        "docx": extract_docx,
        "txt": extract_text,
        "md": extract_text,
    }
    extractor = extractors.get(file_type.lower())
    if not extractor:
        raise ValueError(f"不支持的文件类型: {file_type}")
    return extractor(content)


# ═══════════════════════════════════════════════════════════
# Chunking Engine
# ═══════════════════════════════════════════════════════════

def chunk_text(
    text: str,
    file_id: str,
    chunk_size: int = 500,
    overlap_pct: float = 0.2,
    preserve_paragraphs: bool = True,
) -> List[Chunk]:
    """
    Split text into overlapping chunks.

    Args:
        text: Raw text to split
        file_id: Parent file ID
        chunk_size: Target chunk size in characters
        overlap_pct: Overlap between chunks (0.0–0.5)
        preserve_paragraphs: Try to break at paragraph boundaries

    Returns:
        List of Chunk objects
    """
    import uuid

    if not text.strip():
        return []

    chunks: List[Chunk] = []

    if preserve_paragraphs and "\n\n" in text:
        # Split by paragraphs first
        paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    else:
        # Split by sentences
        paragraphs = _split_sentences(text)

    current_text = ""
    index = 0

    for para in paragraphs:
        # If single paragraph exceeds chunk_size, split it
        if len(para) > chunk_size:
            # Flush current buffer first
            if current_text:
                chunks.append(Chunk(
                    chunk_id=str(uuid.uuid4()),
                    file_id=file_id,
                    content=current_text.strip(),
                    index=index,
                ))
                index += 1
                # Keep overlap from previous chunk
                overlap_text = current_text[-int(len(current_text) * overlap_pct):] if overlap_pct > 0 else ""
                current_text = overlap_text

            # Split long paragraph into sentences
            sentences = _split_sentences(para)
            for sent in sentences:
                if len(current_text) + len(sent) + 1 > chunk_size and current_text:
                    chunks.append(Chunk(
                        chunk_id=str(uuid.uuid4()),
                        file_id=file_id,
                        content=current_text.strip(),
                        index=index,
                    ))
                    index += 1
                    overlap_text = current_text[-int(len(current_text) * overlap_pct):] if overlap_pct > 0 else ""
                    current_text = overlap_text
                current_text += " " + sent if current_text else sent
            continue

        # Normal paragraph that fits in chunk_size
        if len(current_text) + len(para) + 2 > chunk_size and current_text:
            chunks.append(Chunk(
                chunk_id=str(uuid.uuid4()),
                file_id=file_id,
                content=current_text.strip(),
                index=index,
            ))
            index += 1
            overlap_text = current_text[-int(len(current_text) * overlap_pct):] if overlap_pct > 0 else ""
            current_text = overlap_text

        current_text += "\n\n" + para if current_text else para

    # Flush remaining
    if current_text.strip():
        chunks.append(Chunk(
            chunk_id=str(uuid.uuid4()),
            file_id=file_id,
            content=current_text.strip(),
            index=index,
        ))

    return chunks


def _split_sentences(text: str) -> List[str]:
    """Split text into sentences using regex."""
    # Handle Chinese + English sentence boundaries
    sentences = re.split(r'(?<=[。！？\.\!\?])\s*', text)
    return [s for s in sentences if s.strip()]


# ═══════════════════════════════════════════════════════════
# Pipeline Orchestrator
# ═══════════════════════════════════════════════════════════

async def process_file(
    content: bytes,
    file_id: str,
    file_type: str,
    chunk_size: int = 500,
    overlap_pct: float = 0.2,
) -> Dict[str, Any]:
    """
    Full pipeline: extract text → chunk → return chunks for embedding.

    Returns:
        {
            file_id,
            raw_text,
            chunks: [Chunk, ...],
            chunk_count,
            status: "ok" | "error",
            error: str | None,
        }
    """
    try:
        raw_text = extract_text_from_file(content, file_type)
        chunks = chunk_text(raw_text, file_id, chunk_size, overlap_pct)
        return {
            "file_id": file_id,
            "raw_text": raw_text,
            "chunks": chunks,
            "chunk_count": len(chunks),
            "status": "ok",
            "error": None,
        }
    except Exception as e:
        return {
            "file_id": file_id,
            "raw_text": "",
            "chunks": [],
            "chunk_count": 0,
            "status": "error",
            "error": str(e),
        }
